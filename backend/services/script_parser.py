from docx import Document
from docx.shared import RGBColor
from typing import Dict, List, Tuple
import re

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


def rgb_to_hex(rgb: RGBColor) -> str:
    """Convert RGBColor to hex string."""
    if rgb is None:
        return "#000000"
    return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}".upper()


def get_run_color(run) -> str:
    """Extract color from a run, checking both direct color and style."""
    if run.font.color.rgb:
        return rgb_to_hex(run.font.color.rgb)
    
    if run.font.color.theme_color:
        return "#000000"
    
    return "#000000"


def parse_docx_script(filepath: str) -> Dict[str, List[Tuple[int, str]]]:
    """
    Parse a DOCX file and extract lines grouped by color.
    
    Returns:
        Dict mapping color hex codes to list of (line_number, text) tuples
    """
    doc = Document(filepath)
    lines_by_color: Dict[str, List[Tuple[int, str]]] = {}
    line_number = 0
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
            
        line_number += 1
        current_color = None
        current_text = []
        
        for run in para.runs:
            if not run.text.strip():
                continue
                
            color = get_run_color(run)
            
            if current_color is None:
                current_color = color
            
            if color == current_color:
                current_text.append(run.text)
            else:
                if current_text and current_color:
                    text = "".join(current_text).strip()
                    if text:
                        if current_color not in lines_by_color:
                            lines_by_color[current_color] = []
                        lines_by_color[current_color].append((line_number, text))
                
                current_color = color
                current_text = [run.text]
        
        if current_text and current_color:
            text = "".join(current_text).strip()
            if text:
                if current_color not in lines_by_color:
                    lines_by_color[current_color] = []
                lines_by_color[current_color].append((line_number, text))
    
    return lines_by_color


def parse_docx_with_all_lines(filepath: str) -> List[Tuple[int, str, str]]:
    """
    Parse a DOCX file and return all lines with their colors.
    
    Returns:
        List of (line_number, text, color_hex) tuples in order
    """
    doc = Document(filepath)
    all_lines: List[Tuple[int, str, str]] = []
    line_number = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        line_number += 1
        
        dominant_color = "#000000"
        max_length = 0
        
        for run in para.runs:
            if run.text.strip():
                color = get_run_color(run)
                if len(run.text) > max_length:
                    max_length = len(run.text)
                    dominant_color = color
        
        all_lines.append((line_number, text, dominant_color))
    
    return all_lines


def get_unique_colors(filepath: str) -> List[str]:
    """Get list of unique colors found in the document."""
    lines = parse_docx_with_all_lines(filepath)
    colors = list(set(color for _, _, color in lines))
    return sorted(colors)


def _pdf_char_color_to_hex(char: dict) -> str:
    """Convert pdfplumber char non_stroking_color to hex. Default #000000 if missing."""
    color = char.get("non_stroking_color")
    if color is None:
        return "#000000"
    if not isinstance(color, (list, tuple)) or len(color) < 3:
        return "#000000"
    # RGB 0-1 or 0-255
    r, g, b = color[0], color[1], color[2]
    if all(0 <= x <= 1 for x in (r, g, b)):
        r, g, b = int(r * 255), int(g * 255), int(b * 255)
    else:
        r, g, b = int(r), int(g), int(b)
    return f"#{r:02x}{g:02x}{b:02x}".upper()


def parse_pdf_with_all_lines(filepath: str) -> List[Tuple[int, str, str]]:
    """
    Parse a PDF and return lines with (line_number, text, color_hex).
    Groups by vertical position (line); uses character fill color when available.
    """
    if pdfplumber is None:
        raise ImportError("pdfplumber is required for PDF uploads. pip install pdfplumber")
    all_lines: List[Tuple[int, str, str]] = []
    line_number = 0
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            chars = page.chars
            if not chars:
                # Fallback: extract text by line
                text = page.extract_text()
                if text:
                    for raw_line in text.splitlines():
                        line_text = raw_line.strip()
                        if line_text:
                            line_number += 1
                            all_lines.append((line_number, line_text, "#000000"))
                continue
            # Group chars by approximate line (top within 3px)
            lines_by_top: Dict[float, List[dict]] = {}
            for c in chars:
                top = round(c.get("top", 0) / 3) * 3
                if top not in lines_by_top:
                    lines_by_top[top] = []
                lines_by_top[top].append(c)
            for top in sorted(lines_by_top.keys()):
                line_chars = lines_by_top[top]
                line_chars.sort(key=lambda x: (x.get("x0", 0), x.get("x1", 0)))
                # Build text and dominant color for the line
                parts = []
                color_counts: Dict[str, int] = {}
                for c in line_chars:
                    t = c.get("text") or ""
                    parts.append(t)
                    hex_color = _pdf_char_color_to_hex(c)
                    color_counts[hex_color] = color_counts.get(hex_color, 0) + len(t)
                text = "".join(parts).strip()
                if not text:
                    continue
                line_number += 1
                dominant_color = max(color_counts, key=color_counts.get) if color_counts else "#000000"
                all_lines.append((line_number, text, dominant_color))
    return all_lines


def extract_character_names(text: str) -> Tuple[str, str]:
    """
    Try to extract character name from line if it follows pattern like:
    "CHARACTER: dialogue" or "CHARACTER - dialogue"
    
    Returns:
        Tuple of (character_name, dialogue) or (None, original_text)
    """
    patterns = [
        r'^([A-Z][A-Z\s]+):\s*(.+)$',
        r'^([A-Z][A-Z\s]+)\s*-\s*(.+)$',
        r'^\[([^\]]+)\]\s*(.+)$',
    ]
    
    for pattern in patterns:
        match = re.match(pattern, text)
        if match:
            return match.group(1).strip(), match.group(2).strip()
    
    return None, text
